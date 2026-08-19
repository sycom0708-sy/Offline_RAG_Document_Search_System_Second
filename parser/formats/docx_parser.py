"""DOCX 파서 (T1.3) — python-docx 기반 문단/표/이미지 분리 추출."""

from __future__ import annotations

from pathlib import Path

import docx
from docx.document import Document as DocxDocument
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph

from parser.base import BaseParser, DocumentReadError
from parser.schema import ImageData, ParsedDocument, TableData
from parser.utils.headings import body_size_of, clean_heading, is_heading_size
from parser.utils.libreoffice import LibreOfficeError
from parser.utils.render import render_pages

_IMAGE_CONTENT_PREFIX = "image/"


class DocxParser(BaseParser):
    extensions = (".docx",)

    def __init__(self, asset_dir: str | Path | None = None, capture_vector_shapes: bool = False) -> None:
        super().__init__(asset_dir)
        # 벡터 도형 캡처는 LibreOffice 변환을 거쳐 느리므로 기본 OFF.
        self._capture_vector_shapes = capture_vector_shapes

    def _parse(self, path: Path, document: ParsedDocument) -> None:
        try:
            source = docx.Document(str(path))
        except Exception as exc:
            raise DocumentReadError(f"DOCX를 열 수 없습니다: {path.name} ({exc})") from exc

        document.title = self._extract_title(source) or path.stem
        asset_dir = self.asset_dir_for(path)

        paragraph_buffer: list[str] = []
        heading = ""
        blocks = list(self._iter_blocks(source))
        fallback = self._font_size_headings(blocks)
        paragraph_index = -1
        for block in blocks:
            if isinstance(block, Paragraph):
                paragraph_index += 1
                text = block.text.strip()
                if not text:
                    continue
                if self._is_heading(block) or paragraph_index in fallback:
                    # 새 절이 시작된다 — 앞선 문단은 **이전** 절의 제목으로 확정하고
                    # 나서 제목을 갈아 끼운다. 순서를 바꾸면 앞 절 내용이 다음 절
                    # 제목을 달고 나온다.
                    self._flush_text(document, paragraph_buffer, heading)
                    heading = clean_heading(text)
                # 제목 문단도 본문에 그대로 남긴다 — 빼면 검색에서 그 문구를 못 찾는다.
                paragraph_buffer.append(text)
            else:
                # 표를 만나면 앞선 문단을 먼저 확정해 텍스트/표 청크가 섞이지 않게 한다.
                self._flush_text(document, paragraph_buffer, heading)
                table_data = self._read_table(block)
                if table_data is not None:
                    document.chunks.append(
                        self.make_table_chunk(document, table_data, heading=heading)
                    )

        self._flush_text(document, paragraph_buffer, heading)
        self._extract_images(document, source, asset_dir)
        self._capture_drawings(document, path, asset_dir)

    @staticmethod
    def _iter_blocks(source: DocxDocument):
        """본문 요소를 문서 순서대로 순회한다 (문단/표 위치 관계 보존)."""
        body = source.element.body
        for child in body.iterchildren():
            if child.tag == qn("w:p"):
                yield Paragraph(child, source)
            elif child.tag == qn("w:tbl"):
                yield Table(child, source)

    def _flush_text(
        self, document: ParsedDocument, buffer: list[str], heading: str = ""
    ) -> None:
        body = "\n".join(buffer).strip()
        if body:
            document.chunks.append(self.make_text_chunk(document, body, heading=heading))
        buffer.clear()

    @classmethod
    def _font_size_headings(cls, blocks: list) -> set[int]:
        """Heading 스타일이 아예 없는 문서를 위한 **글꼴 크기 폴백** (T10.32).

        `.doc`이 이 경우다 — LibreOffice로 docx 변환을 거쳐도 스타일 이름이
        `MS바탕글`·`Normal`뿐이라 스타일로는 제목을 찾을 수 없다(실측: 9개 전부).
        대신 변환 과정에서 글꼴 크기가 **명시값으로** 박히므로 크기로는 찾을 수 있다.

        🔴 **문서에서 가장 큰 크기만 제목으로 본다.** 문턱(`is_heading_size`)만
        쓰면 본문 10pt 문서에서 12pt 줄이 전부 걸려, 이어달리기 방식이라 진짜 제목
        (16pt)을 뒤따르는 노이즈가 덮어쓴다 — 실측: `6.코치추천서`에서 `코치 추천서`
        뒤에 `20  .    .`·`(사)한국코치협회 귀하`가 12pt로 붙어 있었다. 최대 크기만
        쓰면 같은 8개 문서에서 6개 정확·오탐 0이 된다.

        스타일이 하나라도 있으면 이 폴백을 쓰지 않는다 — 두 기준을 섞으면 제목
        수준이 뒤죽박죽이 된다. 순정 docx는 크기가 테마 기본값이라 애초에
        `None`으로 나와(실측 3개 문서) 이 폴백이 걸리지 않는다.

        🔴 반환값은 **문단 순번**이다. lxml은 같은 노드라도 접근할 때마다 새 프록시를
        내주므로 요소 자체나 `id()`로는 "이 문단이 그 문단인지"를 판별할 수 없다.
        """
        paragraphs = [b for b in blocks if isinstance(b, Paragraph)]
        if any(cls._is_heading(p) for p in paragraphs):
            return set()

        sized: list[tuple[float, str, int]] = []
        for index, paragraph in enumerate(paragraphs):
            text = paragraph.text.strip()
            if not text:
                continue
            size = cls._paragraph_font_size(paragraph)
            if size:
                sized.append((size, text, index))
        if not sized:
            return set()

        body_size = body_size_of([(size, text) for size, text, _ in sized])
        largest = max(size for size, _, _ in sized)
        if not is_heading_size(largest, body_size):
            return set()

        return {
            index for size, text, index in sized if size == largest and clean_heading(text)
        }

    @staticmethod
    def _paragraph_font_size(paragraph: Paragraph) -> float:
        """문단에 쓰인 가장 큰 글꼴 크기(pt). 명시값이 없으면 스타일을 거슬러 올라간다."""
        sizes = [run.font.size.pt for run in paragraph.runs if run.font.size is not None]
        if sizes:
            return max(sizes)

        style = paragraph.style
        while style is not None and style.font.size is None:
            style = getattr(style, "base_style", None)
        if style is not None and style.font.size is not None:
            return style.font.size.pt
        return 0.0

    @staticmethod
    def _is_heading(paragraph: Paragraph) -> bool:
        """Heading/Title 스타일 문단인가 (T10.31).

        `_extract_title()`이 문서 제목을 찾을 때 쓰는 것과 같은 판정이지만,
        이쪽은 **문서 전체가 아니라 지금 지나온 절**을 추적한다.
        """
        style = getattr(paragraph, "style", None)
        name = getattr(style, "name", "") or ""
        return name.startswith("Heading") or name.startswith("Title")

    @staticmethod
    def _read_table(table: Table) -> TableData | None:
        rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
        return TableData.from_rows(rows)

    def _extract_images(self, document: ParsedDocument, source: DocxDocument, asset_dir: Path) -> None:
        for index, part in enumerate(source.part.related_parts.values()):
            content_type = getattr(part, "content_type", "") or ""
            if not content_type.startswith(_IMAGE_CONTENT_PREFIX):
                continue
            try:
                blob = part.blob
            except Exception as exc:
                self.record_error(document, f"이미지 추출 실패({part.partname}): {exc}")
                continue

            ext = Path(str(part.partname)).suffix or ".png"
            out_path = asset_dir / f"img{index:02d}{ext}"
            out_path.write_bytes(blob)
            document.chunks.append(
                self.make_image_chunk(
                    document,
                    ImageData(image_path=str(out_path), origin="extracted"),
                )
            )

    def _capture_drawings(self, document: ParsedDocument, path: Path, asset_dir: Path) -> None:
        if not self._capture_vector_shapes:
            return
        try:
            captured = render_pages(path, asset_dir)
        except LibreOfficeError as exc:
            self.record_error(document, f"벡터 도형 캡처 실패: {exc}")
            return

        for page_no, image_path in sorted(captured.items()):
            document.chunks.append(
                self.make_image_chunk(
                    document,
                    ImageData(
                        image_path=str(image_path),
                        caption=f"{page_no}쪽 도면 캡처",
                        origin="rendered",
                    ),
                    page_or_slide=page_no,
                )
            )

    @staticmethod
    def _extract_title(source: DocxDocument) -> str:
        core_title = (source.core_properties.title or "").strip()
        if core_title:
            return core_title
        for paragraph in source.paragraphs:
            text = paragraph.text.strip()
            if text and paragraph.style is not None and paragraph.style.name.startswith("Heading"):
                return text[:200]
        for paragraph in source.paragraphs:
            if paragraph.text.strip():
                return paragraph.text.strip()[:200]
        return ""
