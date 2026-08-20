"""형식별 파서 파싱 정확도 테스트 (T1.13)."""

from __future__ import annotations

from pathlib import Path

import pytest

from parser import ChunkType, ParseStatus, parse_file
from tests.fixtures.generate_samples import BODY_PARAGRAPH, BODY_TEXT, TABLE_HEADER, TABLE_ROWS

EXPECTED_TABLE_ROWS = [row for row in TABLE_ROWS]


def text_of(document) -> str:
    return "\n".join(c.content for c in document.chunks_of(ChunkType.TEXT))


# --- TXT (T1.4) ---------------------------------------------------------


def test_txt_parses_utf8(sample_txt):
    document = parse_file(sample_txt)
    assert document.status is ParseStatus.OK
    assert BODY_PARAGRAPH in text_of(document)
    assert document.title == BODY_TEXT


@pytest.mark.parametrize("encoding", ["cp949", "euc-kr", "utf-8-sig", "utf-16"])
def test_txt_detects_korean_encodings(tmp_path, encoding):
    path = tmp_path / f"sample_{encoding}.txt"
    path.write_bytes(f"{BODY_TEXT}\n{BODY_PARAGRAPH}".encode(encoding))
    document = parse_file(path)
    assert BODY_PARAGRAPH in text_of(document)


def test_txt_handles_korean_path(tmp_path):
    korean_dir = tmp_path / "한글폴더"
    korean_dir.mkdir()
    path = korean_dir / "한글문서.txt"
    path.write_text(BODY_PARAGRAPH, encoding="cp949")
    document = parse_file(path)
    assert BODY_PARAGRAPH in text_of(document)
    assert document.file_name == "한글문서.txt"


# --- PDF (T1.2) ---------------------------------------------------------


def test_pdf_extracts_text_table_and_images(sample_pdf):
    document = parse_file(sample_pdf)
    assert document.status is ParseStatus.OK
    assert BODY_TEXT in text_of(document)

    tables = document.chunks_of(ChunkType.TABLE)
    assert len(tables) == 1
    assert tables[0].table.header_row == TABLE_HEADER
    assert tables[0].table.rows == EXPECTED_TABLE_ROWS

    images = document.chunks_of(ChunkType.IMAGE)
    assert images, "삽입 이미지가 추출되지 않았습니다"
    assert all(Path(c.image.image_path).is_file() for c in images)


def test_pdf_renders_vector_diagram_page(sample_pdf):
    """벡터 도형은 이미지로 추출되지 않으므로 페이지 렌더링 캡처가 있어야 한다 (TECH 3.1절)."""
    document = parse_file(sample_pdf)
    rendered = [c for c in document.chunks_of(ChunkType.IMAGE) if c.image.origin == "rendered"]
    assert len(rendered) == 1
    assert rendered[0].page_or_slide == 2
    assert Path(rendered[0].image.image_path).is_file()


def test_pdf_records_page_numbers(sample_pdf):
    document = parse_file(sample_pdf)
    assert all(c.page_or_slide is not None for c in document.chunks)


# --- DOCX (T1.3) --------------------------------------------------------


def test_docx_extracts_text_table_and_image(sample_docx):
    document = parse_file(sample_docx)
    assert document.status is ParseStatus.OK
    assert document.title == BODY_TEXT
    assert BODY_PARAGRAPH in text_of(document)

    tables = document.chunks_of(ChunkType.TABLE)
    assert len(tables) == 1
    assert tables[0].table.header_row == TABLE_HEADER
    assert tables[0].table.rows == EXPECTED_TABLE_ROWS

    images = document.chunks_of(ChunkType.IMAGE)
    assert len(images) == 1
    assert Path(images[0].image.image_path).is_file()


def test_docx_keeps_paragraphs_around_table_separate(sample_docx):
    """표 앞뒤 문단이 하나의 청크로 뭉치지 않아야 한다."""
    document = parse_file(sample_docx)
    texts = document.chunks_of(ChunkType.TEXT)
    assert len(texts) == 2
    assert "표 아래 문단입니다." == texts[1].content


# --- XLSX (T1.5) --------------------------------------------------------


def test_xlsx_extracts_each_sheet_as_table(sample_xlsx):
    document = parse_file(sample_xlsx)
    assert document.status is ParseStatus.OK

    tables = document.chunks_of(ChunkType.TABLE)
    assert len(tables) == 2
    assert tables[0].table.caption == "사양표"
    assert tables[0].table.header_row == TABLE_HEADER
    assert tables[0].table.rows == EXPECTED_TABLE_ROWS
    assert tables[0].page_or_slide == 1


def test_xlsx_extracts_embedded_image(sample_xlsx):
    document = parse_file(sample_xlsx)
    images = document.chunks_of(ChunkType.IMAGE)
    assert len(images) == 1
    assert Path(images[0].image.image_path).is_file()


# --- PPTX (T1.6) --------------------------------------------------------


def test_pptx_extracts_per_slide_content(sample_pptx):
    document = parse_file(sample_pptx)
    assert document.status is ParseStatus.OK
    assert BODY_TEXT in text_of(document)

    tables = document.chunks_of(ChunkType.TABLE)
    assert len(tables) == 1
    assert tables[0].page_or_slide == 2
    assert tables[0].table.header_row == TABLE_HEADER
    assert tables[0].table.rows == EXPECTED_TABLE_ROWS

    images = document.chunks_of(ChunkType.IMAGE)
    assert len(images) == 1
    assert images[0].page_or_slide == 2


# --- HWPX (T1.8) --------------------------------------------------------


def test_hwpx_extracts_text_table_and_image(sample_hwpx):
    document = parse_file(sample_hwpx)
    assert document.status is ParseStatus.OK
    assert BODY_PARAGRAPH in text_of(document)

    tables = document.chunks_of(ChunkType.TABLE)
    assert len(tables) == 1
    assert tables[0].table.header_row == TABLE_HEADER
    assert tables[0].table.rows == EXPECTED_TABLE_ROWS

    images = document.chunks_of(ChunkType.IMAGE)
    assert len(images) == 1
    assert Path(images[0].image.image_path).is_file()


def test_hwpx_keeps_paragraphs_around_table_separate(sample_hwpx):
    document = parse_file(sample_hwpx)
    texts = document.chunks_of(ChunkType.TEXT)
    assert len(texts) == 2
    assert texts[1].content == "표 아래 문단입니다."


def test_hwpx_rejects_non_zip(tmp_path):
    from parser import DocumentReadError

    path = tmp_path / "broken.hwpx"
    path.write_bytes(b"not a zip file")
    with pytest.raises(DocumentReadError):
        parse_file(path)


# --- 공통 규칙 (T1.10 / T1.11 / T1.12) -----------------------------------

ALL_SAMPLE_KEYS = ["sample.txt", "sample.pdf", "sample.docx", "sample.xlsx", "sample.pptx", "sample.hwpx"]


@pytest.mark.parametrize("sample_key", ALL_SAMPLE_KEYS)
def test_every_parser_emits_common_schema(samples, sample_key):
    document = parse_file(samples[sample_key])
    assert document.chunks, f"{sample_key}에서 청크가 추출되지 않았습니다"
    for chunk in document.chunks:
        assert chunk.doc_id == document.doc_id
        assert chunk.file_name == document.file_name
        assert chunk.source_hash == document.source_hash
        assert chunk.source_mtime == document.source_mtime
        assert chunk.created_at
        assert chunk.embedding_vector is None  # Phase 3에서 채워진다


@pytest.mark.parametrize("sample_key", ALL_SAMPLE_KEYS)
def test_chunk_ids_are_unique(samples, sample_key):
    document = parse_file(samples[sample_key])
    ids = [c.chunk_id for c in document.chunks]
    assert len(ids) == len(set(ids))


@pytest.mark.parametrize("sample_key", ALL_SAMPLE_KEYS)
def test_tables_are_never_merged_into_text_chunks(samples, sample_key):
    """표 내용이 텍스트 청크에 섞이면 행·열 구조가 소실된다 (TECH 3.1절)."""
    document = parse_file(samples[sample_key])
    for chunk in document.chunks_of(ChunkType.TEXT):
        assert "8GB | 16GB" not in chunk.content


@pytest.mark.parametrize("sample_key", ALL_SAMPLE_KEYS)
def test_table_chunks_carry_structured_data(samples, sample_key):
    document = parse_file(samples[sample_key])
    for chunk in document.chunks_of(ChunkType.TABLE):
        assert chunk.table is not None
        assert isinstance(chunk.table.rows, list)
        # 캡션/헤더는 FTS5 키워드 가중에 쓰인다 (TECH 4.3절)
        assert chunk.keywords


@pytest.mark.parametrize("sample_key", ALL_SAMPLE_KEYS)
def test_image_chunks_point_to_existing_files(samples, sample_key):
    document = parse_file(samples[sample_key])
    for chunk in document.chunks_of(ChunkType.IMAGE):
        assert chunk.image is not None
        assert Path(chunk.image.image_path).is_file()
        assert chunk.image.origin in {"extracted", "rendered"}


# --- T10.31: 절 제목(heading) ------------------------------------------------


class _FakeSpan(dict):
    pass


def _page_stub(lines):
    """`page.get_text("dict")` 모양의 최소 스텁. (글꼴크기, 텍스트) 목록을 받는다."""

    class _Page:
        def get_text(self, kind):
            assert kind == "dict"
            return {
                "blocks": [
                    {"lines": [{"spans": [{"size": size, "text": text}]} for size, text in lines]}
                ]
            }

    return _Page()


class TestPdfPageHeading:
    """🔴 텍스트 패턴이 아니라 **글꼴 크기**로 제목을 찾는다.

    이 코퍼스의 PDF는 페이지 전체가 줄바꿈 없는 한 덩어리로 추출돼(실측:
    text 청크 972개 중 952개가 단일 줄) "짧은 줄"이라는 단서가 없다.
    """

    def _heading(self, lines):
        from parser.formats.pdf_parser import PdfParser

        return PdfParser()._page_heading(_page_stub(lines))

    def test_largest_font_line_becomes_the_heading(self):
        """실측 재현: AICA 안내서 3쪽은 제목 20pt / 본문 14pt였다."""
        assert self._heading([
            (12.0, "3 / 20 정도 · 혁신 · 협업"),
            (20.0, "1-1. AICA 취득 절차"),
            (14.0, "인증 : AICA(AIL / AIU / AIC / AIS)"),
        ]) == "1-1. AICA 취득 절차"

    def test_uniform_font_page_has_no_heading(self):
        """글꼴이 균일하면 제목을 가릴 근거가 없다 — 아무 줄이나 올리면 노이즈다."""
        assert self._heading([
            (12.0, "첫 문단입니다"),
            (12.0, "둘째 문단입니다"),
        ]) == ""

    def test_slightly_larger_font_is_not_a_heading(self):
        """본문과 충분히 구분되지 않으면(비율 미달) 제목으로 보지 않는다."""
        assert self._heading([
            (12.5, "조금 큰 줄"),
            (12.0, "본문입니다"),
        ]) == ""

    def test_long_line_is_rejected_not_truncated(self):
        """🔴 잘라 쓰면 본문 앞부분이 제목으로 둔갑한다 — 통째로 버려야 한다.

        실측: PBV01 문서에서 54자짜리 본문 문단이 그 페이지 최대 글꼴이라
        제목으로 잡혔다.
        """
        assert self._heading([
            (20.0, "개발시adb install이나 개발툴에서 설치를 하게 되면 일반앱으로 설치가 되어 기존의 방법대로"),
            (12.0, "본문"),
        ]) == ""

    def test_only_first_line_of_largest_font_is_used(self):
        """같은 크기 줄을 전부 이으면 목차 페이지가 통째로 붙는다(실측: DTG 문서)."""
        assert self._heading([
            (20.0, "1. 개요"),
            (20.0, "2. 참고"),
            (12.0, "본문"),
        ]) == "1. 개요"


class TestPdfTableHeadingSplit:
    """표 안에 섞인 절 제목 행으로 표를 쪼갠다.

    실측: 기아차 앱미터기 결제 프로토콜정의서(PDF)는 4.1~4.6 여러 절의 표를
    하나의 연속된 표로 그리면서, 절 제목("4.2 업무개시 수신 응답 ...")을
    표의 한 행(첫 칸만 채움)으로 끼워 넣는다. 검색어가 그 제목에 매치되면
    앞 절의 표까지 결과에 함께 나오던 문제(사용자 보고)의 원인이다.
    """

    def _split(self, rows):
        from parser.formats.pdf_parser import PdfParser

        return PdfParser._split_table_rows_on_heading(rows)

    def test_no_heading_row_stays_a_single_segment(self):
        """제목 행이 없으면 기존과 동일하게 표 전체가 구간 하나다."""
        rows = [["Header", "1", "STX"], ["", "", "OP CODE"]]
        segments = self._split(rows)
        assert len(segments) == 1
        assert segments[0] == ("", rows)

    def test_embedded_heading_row_splits_the_table(self):
        """실측 재현: 4.1 표 꼬리 뒤에 '4.2 ...' 행이 오면 그 지점에서 쪼갠다."""
        tail_of_4_1 = ["", "총 Byte 길이", "", "", "", ""]
        heading_row = ["4.2 업무개시 수신 응답 (결제기 → 앱미터기)", "", "", "", "", ""]
        head_of_4_2 = ["Header", "1", "STX", "", "", ""]

        segments = self._split([tail_of_4_1, heading_row, head_of_4_2])

        assert len(segments) == 2
        assert segments[0] == ("", [tail_of_4_1])
        assert segments[1] == ("4.2 업무개시 수신 응답 (결제기 → 앱미터기)", [head_of_4_2])

    def test_multiple_heading_rows_split_into_multiple_segments(self):
        rows = [
            ["Header", "1", "STX"],
            ["4.2 업무개시 수신 응답 (결제기 → 앱미터기)", "", "", "", "", ""],
            ["Header", "1", "STX"],
            ["4.4 결제 결과 응답 (앱미터기 → 결제기)", "", "", "", "", ""],
            ["Header", "1", "STX"],
        ]
        segments = self._split(rows)
        assert [heading for heading, _ in segments] == [
            "",
            "4.2 업무개시 수신 응답 (결제기 → 앱미터기)",
            "4.4 결제 결과 응답 (앱미터기 → 결제기)",
        ]
        assert [len(rows) for _, rows in segments] == [1, 1, 1]

    def test_row_with_other_filled_cells_is_not_a_heading(self):
        """첫 칸이 번호 패턴이어도 다른 칸이 채워져 있으면 데이터 행이다."""
        rows = [["4.2", "업무개시", "수신 응답"], ["Header", "1", "STX"]]
        segments = self._split(rows)
        assert len(segments) == 1
        assert segments[0] == ("", rows)

    def test_number_without_following_text_is_not_a_heading(self):
        """번호만 있고 뒤에 텍스트가 없으면(예: 순번 칸) 제목으로 보지 않는다."""
        rows = [["4.2", "", "", "", "", ""], ["Header", "1", "STX"]]
        segments = self._split(rows)
        assert len(segments) == 1
        assert segments[0] == ("", rows)

    def test_overlong_first_cell_is_dropped_not_split(self):
        """`clean_heading()`의 길이 상한을 그대로 따른다 — 너무 길면 제목 후보가 아니다."""
        long_text = "4.9 " + "가" * 40
        rows = [[long_text, "", "", "", "", ""], ["Header", "1", "STX"]]
        segments = self._split(rows)
        assert len(segments) == 1
        assert segments[0] == ("", rows)


def test_heading_survives_store_and_search(tmp_path):
    """파서가 뽑은 제목이 DB를 거쳐 검색 결과까지 그대로 도달해야 한다."""
    from indexer.fts5.schema import connect
    from indexer.fts5.search import search as keyword_search
    from indexer.fts5.store import store_document
    from parser.schema import Chunk, ChunkType, ParsedDocument

    conn = connect(":memory:")
    document = ParsedDocument(doc_id="d1", file_path="x", file_name="안내서.pdf", title="t")
    document.chunks.append(
        Chunk(
            chunk_id="c1", doc_id="d1", file_path="x", file_name="안내서.pdf",
            type=ChunkType.TEXT, page_or_slide=3,
            content="인증 절차는 다음과 같습니다",
            heading="1-1. AICA 취득 절차",
        )
    )
    store_document(conn, document)

    results = keyword_search(conn, "인증")
    assert results[0].heading == "1-1. AICA 취득 절차"


def test_heading_is_not_searchable(tmp_path):
    """🔴 제목은 표시 전용이다 — 색인에 넣으면 제목이 걸린 문서의 모든 청크가
    결과에 끼는 T10.6(파일명 매치)과 같은 일이 생긴다."""
    from indexer.fts5.schema import connect
    from indexer.fts5.search import search as keyword_search
    from indexer.fts5.store import store_document
    from parser.schema import Chunk, ChunkType, ParsedDocument

    conn = connect(":memory:")
    document = ParsedDocument(doc_id="d1", file_path="x", file_name="문서.pdf", title="t")
    document.chunks.append(
        Chunk(
            chunk_id="c1", doc_id="d1", file_path="x", file_name="문서.pdf",
            type=ChunkType.TEXT, page_or_slide=1,
            content="본문에는 그 단어가 없다",
            heading="희귀단어제목",
        )
    )
    store_document(conn, document)

    assert keyword_search(conn, "희귀단어제목") == []


def test_heading_column_is_added_to_an_existing_db(tmp_path):
    """구버전 DB에 컬럼만 더한다 — `chunks`는 원문이라 지우면 복구할 수 없다."""
    import sqlite3

    from indexer.fts5.schema import connect

    db = tmp_path / "old.sqlite3"
    conn = connect(db)
    conn.execute(
        "INSERT INTO documents(doc_id, file_path, file_name, status, indexed_at)"
        " VALUES ('d1','x','f.pdf','ok','2026-01-01')"
    )
    conn.execute(
        "INSERT INTO chunks(chunk_id, doc_id, file_path, file_name, type,"
        " content, created_at) VALUES ('c1','d1','x','f.pdf','text','내용','2026-01-01')"
    )
    conn.commit()
    conn.close()

    # 구버전을 흉내내기 위해 컬럼을 떼어낼 수는 없으므로, 재연결이 기존 행을
    # 보존하는지(= 테이블을 다시 만들지 않는지)를 확인한다.
    conn = connect(db)
    assert conn.execute("SELECT COUNT(*) FROM chunks").fetchone()[0] == 1
    assert conn.execute("SELECT heading FROM chunks WHERE chunk_id='c1'").fetchone()[0] == ""
    conn.close()


# ---------------------------------------------------------------------------
# T10.32 — 제목 표시를 나머지 포맷으로 넓힌다 (docx 버그 수정 · xlsx · hwp · hwpx)
# ---------------------------------------------------------------------------


class TestDocxHeadingTracking:
    """🔴 T10.31이 `_is_heading()`을 만들어 놓고 **어디에서도 부르지 않았다**.

    그래서 docx는 Heading 스타일이 멀쩡히 있어도 제목이 늘 빈 문자열이었다
    (실측: `결제기_기아차...프로토콜정의서.docx`는 Heading 문단 15개인데 제목 0개).
    `.doc`·`.rtf`도 LibreOffice 변환 뒤 이 파서에 위임하므로 함께 막혀 있었다.
    """

    def _parse(self, tmp_path):
        import docx

        from parser.formats.docx_parser import DocxParser

        source = docx.Document()
        source.add_paragraph("표지 문단")              # 제목 이전 → 제목 없음
        source.add_paragraph("통신 메시지 구조 정의", style="Heading 1")
        source.add_paragraph("본문 첫 절")
        source.add_paragraph("데이터형식", style="Heading 2")
        source.add_paragraph("본문 둘째 절")
        path = tmp_path / "sample.docx"
        source.save(path)
        return DocxParser(asset_dir=tmp_path / "assets").parse(path)

    def test_headings_are_tracked_across_sections(self, tmp_path):
        document = self._parse(tmp_path)
        texts = [c for c in document.chunks if c.type.value == "text"]
        assert [c.heading for c in texts] == ["", "통신 메시지 구조 정의", "데이터형식"]

    def test_heading_paragraph_stays_in_the_body(self, tmp_path):
        """제목 문단을 본문에서 빼면 그 문구로는 검색이 안 된다."""
        document = self._parse(tmp_path)
        body = "\n".join(c.content for c in document.chunks if c.type.value == "text")
        assert "통신 메시지 구조 정의" in body

    def test_preceding_text_keeps_the_previous_heading(self, tmp_path):
        """제목을 만나면 **앞선 문단부터** 확정한다 — 순서를 바꾸면 앞 절 내용이
        다음 절 제목을 달고 나온다."""
        document = self._parse(tmp_path)
        first = next(c for c in document.chunks if c.type.value == "text")
        assert first.content.startswith("표지 문단")
        assert first.heading == ""


class TestXlsxSheetHeading:
    """시트 1행의 제목 칸을 쓰되, 시트명과 겹치면 비운다 [사용자 확정]."""

    def _heading(self, rows, sheet_title):
        from parser.formats.xlsx_parser import XlsxParser
        from parser.schema import TableData

        return XlsxParser._sheet_heading(TableData.from_rows(rows, caption=sheet_title), sheet_title)

    def test_single_filled_cell_in_first_row_is_the_heading(self):
        assert self._heading([["개발 관련 파일 목록", "", ""], ["a", "b", "c"]], "Sheet1") == "개발 관련 파일 목록"

    def test_heading_equal_to_sheet_name_is_dropped(self):
        """표 카드가 이미 시트명을 위치로 보여준다 — 같은 문자열을 두 줄로 띄우지 않는다."""
        assert self._heading([["2.주행 시작", "", ""], ["a", "b", "c"]], "2.주행 시작") == ""

    def test_whitespace_only_difference_still_counts_as_equal(self):
        assert self._heading([["2. 주행 시작", ""], ["a", "b"]], "2.주행 시작") == ""

    def test_column_header_row_is_not_a_heading(self):
        """여러 칸이 차 있으면 제목이 아니라 열 머리글이고, 그건 표가 이미 보여준다."""
        assert self._heading([["번호", "이름", "비고"], ["1", "김", ""]], "Sheet1") == ""


class TestHwpxHeadingByFontSize:
    """HWPX는 `outlineLevel`이 없어(실측 2문서) 글꼴 크기로 판별한다.

    크기는 문단이 아니라 `header.xml`의 `charPr`에 있어 두 파일을 맞물려 읽는다.
    """

    HEADER = (
        '<hh:head xmlns:hh="http://www.hancom.co.kr/hwpml/2011/head">'
        '<hh:charPr id="0" height="1000"/><hh:charPr id="1" height="1600"/>'
        "</hh:head>"
    )

    def _section(self, paragraphs):
        runs = "".join(
            f'<hp:p><hp:run charPrIDRef="{ref}"><hp:t>{text}</hp:t></hp:run></hp:p>'
            for ref, text in paragraphs
        )
        return (
            '<hs:sec xmlns:hs="http://www.hancom.co.kr/hwpml/2011/section"'
            ' xmlns:hp="http://www.hancom.co.kr/hwpml/2011/paragraph">' + runs + "</hs:sec>"
        )

    def _parse(self, tmp_path, paragraphs):
        import zipfile

        from parser.formats.hwpx_parser import HwpxParser

        path = tmp_path / "sample.hwpx"
        with zipfile.ZipFile(path, "w") as archive:
            archive.writestr("Contents/header.xml", self.HEADER)
            archive.writestr("Contents/section0.xml", self._section(paragraphs))
        return HwpxParser(asset_dir=tmp_path / "assets").parse(path)

    def test_char_heights_are_read_from_header_xml(self, tmp_path):
        """크기가 본문 파일에 없다는 점이 PDF와 다르다 — header.xml을 못 읽으면 제목도 없다."""
        import zipfile

        from parser.formats.hwpx_parser import HwpxParser

        path = tmp_path / "h.hwpx"
        with zipfile.ZipFile(path, "w") as archive:
            archive.writestr("Contents/header.xml", self.HEADER)
        with zipfile.ZipFile(path) as archive:
            assert HwpxParser._char_heights(archive) == {"0": 1000.0, "1": 1600.0}

    def test_missing_header_xml_is_not_an_error(self, tmp_path):
        """제목은 부가 정보다 — 없다고 본문 추출을 막으면 안 된다."""
        import zipfile

        from parser.formats.hwpx_parser import HwpxParser

        path = tmp_path / "h.hwpx"
        with zipfile.ZipFile(path, "w") as archive:
            archive.writestr("Contents/section0.xml", self._section([("0", "본문")]))
        with zipfile.ZipFile(path) as archive:
            assert HwpxParser._char_heights(archive) == {}

    def test_heading_applies_to_following_text(self, tmp_path):
        document = self._parse(
            tmp_path,
            [("0", "본문이 길게 이어지는 문단입니다 " * 3), ("1", "2. 설치 방법"), ("0", "설치 절차 본문")],
        )
        texts = [c for c in document.chunks if c.type.value == "text"]
        assert [c.heading for c in texts] == ["", "2. 설치 방법"]

    def test_uniform_font_document_has_no_heading(self, tmp_path):
        document = self._parse(tmp_path, [("0", "첫 문단"), ("0", "둘째 문단")])
        assert all(c.heading == "" for c in document.chunks)


class TestHwpHeadingByFontSize:
    """HWP는 pyhwp 중간 XML에서 `Text[charshape-id]` → `CharShape[basesize]`로 크기를 푼다."""

    XML = (
        "<HwpDoc>"
        '<CharShape basesize="900"/><CharShape basesize="1600"/>'
        "<BodyText>"
        '<Paragraph><Text charshape-id="0">본문 문단</Text></Paragraph>'
        '<Paragraph><Text charshape-id="1">2. 설치 방법</Text></Paragraph>'
        "</BodyText></HwpDoc>"
    )

    def _root(self):
        from xml.etree import ElementTree

        return ElementTree.fromstring(self.XML)

    def test_char_sizes_follow_definition_order(self):
        from parser.formats.hwp_parser import HwpParser

        assert HwpParser._char_sizes(self._root()) == [900.0, 1600.0]

    def test_paragraph_size_uses_the_largest_run(self):
        from parser.formats.hwp_parser import HwpParser

        root = self._root()
        sizes = HwpParser._char_sizes(root)
        paragraphs = list(root.iter("Paragraph"))
        assert HwpParser._paragraph_size(paragraphs[0], sizes) == 900.0
        assert HwpParser._paragraph_size(paragraphs[1], sizes) == 1600.0

    def test_body_size_is_weighted_by_length(self):
        """짧은 제목이 개수로 많아도 본문 크기를 빼앗지 않는다."""
        from parser.utils.headings import body_size_of

        assert body_size_of([(1600.0, "제목"), (1600.0, "제목2"), (900.0, "본문 " * 30)]) == 900.0

    def test_table_cell_text_is_not_a_heading_candidate(self):
        """표 안 문단은 별도 청크로 빠지므로 제목 후보가 아니다 (Phase 1 결정).

        실측: 리눅스마스터 기출문제(hwp)의 `1과목 : 리눅스 운영 및 관리`가 바로 이
        경우였다 — 눈에는 절 제목이지만 표 셀 안에 들어 있다.
        """
        from xml.etree import ElementTree

        from parser.formats.hwp_parser import HwpParser

        root = ElementTree.fromstring(
            "<HwpDoc>"
            '<CharShape basesize="900"/><CharShape basesize="1600"/>'
            '<Paragraph><TableControl><Paragraph>'
            '<Text charshape-id="1">1과목 : 리눅스 운영 및 관리</Text>'
            "</Paragraph></TableControl></Paragraph></HwpDoc>"
        )
        sizes = HwpParser._char_sizes(root)
        outer = next(iter(root.iter("Paragraph")))
        assert HwpParser._paragraph_size(outer, sizes) == 0.0


class TestDocxFontSizeFallback:
    """Heading 스타일이 없는 문서(=`.doc` 변환본)를 글꼴 크기로 구한다 (T10.32).

    실측: `.doc` 9개 전부가 변환 후에도 스타일이 `MS바탕글`·`Normal`뿐이라
    스타일로는 제목이 0개였다. 반면 변환이 크기를 명시값으로 박아줘 크기로는 찾힌다.
    """

    def _build(self, tmp_path, rows, styled=False):
        import docx
        from docx.shared import Pt

        source = docx.Document()
        if styled:
            source.add_paragraph("스타일 제목", style="Heading 1")
        for size, text in rows:
            paragraph = source.add_paragraph()
            run = paragraph.add_run(text)
            run.font.size = Pt(size)
        path = tmp_path / "sample.docx"
        source.save(path)
        return path

    def _headings(self, tmp_path, rows, styled=False):
        from parser.formats.docx_parser import DocxParser

        path = self._build(tmp_path, rows, styled)
        document = DocxParser(asset_dir=tmp_path / "assets").parse(path)
        return [c.heading for c in document.chunks if c.type.value == "text"]

    def test_largest_font_paragraph_becomes_the_heading(self, tmp_path):
        headings = self._headings(tmp_path, [
            (16.0, "고객 추천서"),
            (10.0, "본문이 이어집니다 " * 5),
        ])
        assert headings == ["고객 추천서"]

    def test_only_the_largest_size_is_used(self, tmp_path):
        """🔴 문턱만 쓰면 12pt 노이즈가 16pt 진짜 제목을 덮어쓴다.

        실측: `6.코치추천서(KAC)`에서 `코치 추천서`(16pt) 뒤에 `20  .    .`과
        `(사)한국코치협회 귀하`가 12pt로 붙어 있었다.
        """
        headings = self._headings(tmp_path, [
            (16.0, "코치 추천서"),
            (10.0, "본문이 이어집니다 " * 5),
            (12.0, "(사)한국코치협회 귀하"),
            (10.0, "이어지는 본문입니다 " * 5),
        ])
        assert set(headings) == {"코치 추천서"}

    def test_uniform_font_document_has_no_heading(self, tmp_path):
        assert self._headings(tmp_path, [
            (10.0, "첫 문단입니다 " * 5),
            (10.0, "둘째 문단입니다 " * 5),
        ]) == [""]

    def test_style_based_document_ignores_the_fallback(self, tmp_path):
        """두 기준을 섞으면 제목 수준이 뒤죽박죽이 된다 — 스타일이 있으면 그것만 쓴다."""
        headings = self._headings(tmp_path, [(28.0, "아주 큰 본문 줄")], styled=True)
        assert headings == ["스타일 제목"]
